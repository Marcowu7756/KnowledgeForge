from pathlib import Path

from docx import Document

fix = Path(r"D:\KnowledgeForge\data\inbox\_fixtures")
fix.mkdir(parents=True, exist_ok=True)

pdf_path = fix / "phase1_sample.pdf"
pdf_path.write_bytes(
    b"""%PDF-1.4
1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj
2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj
3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources<< /Font<< /F1 5 0 R >> >> >>endobj
4 0 obj<< /Length 78 >>stream
BT /F1 12 Tf 50 750 Td (PAILE Phase1 sample: F = m * a. Force causes acceleration.) Tj ET
endstream
endobj
5 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000266 00000 n 
0000000395 00000 n 
trailer<< /Size 6 /Root 1 0 R >>
startxref
472
%%EOF
"""
)

doc_path = fix / "phase1_sample.docx"
doc = Document()
doc.add_heading("PAILE Phase1 DOCX Sample", level=1)
doc.add_paragraph("Core idea: Force equals mass times acceleration (F = ma).")
doc.add_paragraph("This is a read-only fixture for KnowledgeForge ingest tests.")
doc.save(doc_path)

print("wrote", pdf_path)
print("wrote", doc_path)
