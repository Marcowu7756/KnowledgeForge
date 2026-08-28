"""Non-Cartesian source → settle → express matrix helpers.

One express path per signal source (not N×M). LLM / ASR / OCR / network
are mocked so the matrix exercises wiring offline.
"""

from __future__ import annotations

import json
import traceback
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable, Literal
from unittest.mock import patch

from app.models import IngestedSource, KnowledgeUnit

ExpressKind = Literal[
    "animate_fast",
    "ko_animate",
    "ko_narrate_mock",
    "dry_gate",
    "family_view",
]

Status = Literal["pass", "fail", "skip"]


@dataclass
class RowResult:
    source_id: str
    acquire: str
    settle: str
    express: str
    status: Status = "fail"
    detail: str = ""
    card_path: str = ""
    package_dir: str = ""
    gif_path: str = ""
    error: str = ""
    traceback: str = ""


def fake_unit(*, title: str, source_type: str, source: str, url: str | None = None) -> KnowledgeUnit:
    return KnowledgeUnit(
        title=title,
        source=source,
        type=source_type,  # type: ignore[arg-type]
        url=url,
        summary=f"{title} — integration matrix settle summary.",
        concepts=["MatrixConcept", "SignalSource", title[:12] or "Topic"],
        key_points=["Observe signal", "Settle as KO", "Express one path"],
        mechanisms=["signal → settle → express"],
        relationships=["signal → knowledge"],
        timeline=["2026-08-28 — matrix row"],
        tags=["integration-matrix", source_type],
    )


def _patch_compress(unit: KnowledgeUnit):
    return patch("app.pipeline.compress", return_value=unit)


def _write_minimal_pdf(path: Path, text: str = "PAILE matrix PDF sample F=ma") -> Path:
    # Minimal valid-enough PDF for pypdf/pdfminer style extractors.
    stream = f"BT /F1 12 Tf 50 750 Td ({text}) Tj ET"
    content = stream.encode("latin-1", errors="replace")
    objects = [
        b"1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj\n",
        b"2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj\n",
        (
            b"3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Contents 4 0 R /Resources<< /Font<< /F1 5 0 R >> >> >>endobj\n"
        ),
        f"4 0 obj<< /Length {len(content)} >>stream\n".encode()
        + content
        + b"\nendstream\nendobj\n",
        b"5 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj\n",
    ]
    body = b"".join(objects)
    xref_positions = []
    cursor = len(b"%PDF-1.4\n")
    out = [b"%PDF-1.4\n"]
    for obj in objects:
        xref_positions.append(cursor)
        out.append(obj)
        cursor += len(obj)
    xref_start = cursor
    xref = [b"xref\n0 6\n0000000000 65535 f \n"]
    for pos in xref_positions:
        xref.append(f"{pos:010d} 00000 n \n".encode())
    trailer = (
        b"trailer<< /Size 6 /Root 1 0 R >>\n"
        b"startxref\n"
        + str(xref_start).encode()
        + b"\n%%EOF\n"
    )
    path.write_bytes(b"".join(out + xref + [trailer]))
    return path


def _write_minimal_png(path: Path) -> Path:
    from PIL import Image, ImageDraw

    img = Image.new("RGB", (320, 120), color=(24, 32, 48))
    draw = ImageDraw.Draw(img)
    draw.text((16, 44), "KF matrix OCR", fill=(241, 245, 249))
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)
    return path


def _write_minimal_docx(path: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_heading("KF matrix DOCX", level=1)
    doc.add_paragraph("Force equals mass times acceleration. Integration matrix row.")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def _write_minimal_wav(path: Path) -> Path:
    # Tiny RIFF header; ASR is mocked so payload is unused.
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(b"RIFF$$$$WAVEfmt " + b"\x00" * 40)
    return path


def settle_from_card(card: Path, *, packages: Path, animate: bool, narrate: bool) -> Any:
    from app.harness.pipeline import compile_knowledge

    return compile_knowledge(
        str(card),
        from_card=True,
        animate=animate,
        narrate=narrate,
        fast=True,
        index=False,
        package_id=f"matrix_{card.stem}_{datetime.now(timezone.utc).strftime('%H%M%S%f')}",
    )


def express_animate_fast(card: Path, dest: Path) -> Any:
    from app.expression.animate_engine import animate_from_card

    return animate_from_card(
        card,
        dest_dir=dest,
        fast=True,
        renderer="pillow",
    )


def run_row_txt(work: Path, packages: Path) -> RowResult:
    row = RowResult(
        source_id="txt",
        acquire="ingest_file",
        settle="pipeline.run_file→card",
        express="animate_fast(pillow)",
    )
    try:
        path = work / "note.txt"
        cards = work / "cards"
        cards.mkdir(parents=True, exist_ok=True)
        path.write_text("第一段关于美债。\n\n第二段：收益率与美元。", encoding="utf-8")
        unit = fake_unit(title="Matrix TXT", source_type="txt", source=str(path))
        from app.pipeline import run_file

        with _patch_compress(unit):
            result = run_file(path, dest_dir=cards, index=False)
        row.card_path = str(result.markdown_path)
        anim = express_animate_fast(result.markdown_path, work / "express_txt")
        row.gif_path = str(anim.gif_path)
        assert anim.gif_path.is_file() and anim.gif_path.stat().st_size > 50
        row.status = "pass"
        row.detail = f"renderer={anim.renderer}"
    except Exception as exc:  # noqa: BLE001
        row.status = "fail"
        row.error = str(exc)
        row.traceback = traceback.format_exc()
    return row


def run_row_md(work: Path, packages: Path) -> RowResult:
    row = RowResult(
        source_id="md",
        acquire="ingest_file",
        settle="run_file→card",
        express="ko_animate",
    )
    try:
        path = work / "note.md"
        cards = work / "cards"
        cards.mkdir(parents=True, exist_ok=True)
        path.write_text("# Matrix MD\n\n正文：石油美元机制简述。", encoding="utf-8")
        unit = fake_unit(title="Matrix MD", source_type="md", source=str(path))
        from app.pipeline import run_file

        with _patch_compress(unit):
            result = run_file(path, dest_dir=cards, index=False)
        row.card_path = str(result.markdown_path)
        compiled = settle_from_card(
            result.markdown_path, packages=packages, animate=True, narrate=False
        )
        row.package_dir = str(compiled.package_dir)
        assert compiled.animation_gif and compiled.animation_gif.is_file()
        row.gif_path = str(compiled.animation_gif)
        row.status = "pass"
        row.detail = f"ko={compiled.knowledge_object_path}"
    except Exception as exc:  # noqa: BLE001
        row.status = "fail"
        row.error = str(exc)
        row.traceback = traceback.format_exc()
    return row


def run_row_pdf(work: Path, packages: Path) -> RowResult:
    row = RowResult(
        source_id="pdf",
        acquire="ingest_file/pdf",
        settle="run_file→card",
        express="ko_animate",
    )
    try:
        path = _write_minimal_pdf(work / "sample.pdf")
        cards = work / "cards"
        cards.mkdir(parents=True, exist_ok=True)
        unit = fake_unit(title="Matrix PDF", source_type="pdf", source=str(path))
        from app.pipeline import run_file

        with _patch_compress(unit):
            result = run_file(path, dest_dir=cards, index=False)
        row.card_path = str(result.markdown_path)
        compiled = settle_from_card(
            result.markdown_path, packages=packages, animate=True, narrate=False
        )
        row.package_dir = str(compiled.package_dir)
        assert compiled.animation_gif and compiled.animation_gif.is_file()
        row.gif_path = str(compiled.animation_gif)
        row.status = "pass"
    except Exception as exc:  # noqa: BLE001
        row.status = "fail"
        row.error = str(exc)
        row.traceback = traceback.format_exc()
    return row


def run_row_docx(work: Path, packages: Path) -> RowResult:
    row = RowResult(
        source_id="docx",
        acquire="ingest_file/docx",
        settle="run_file→card",
        express="animate_fast(pillow)",
    )
    try:
        path = _write_minimal_docx(work / "sample.docx")
        cards = work / "cards"
        cards.mkdir(parents=True, exist_ok=True)
        unit = fake_unit(title="Matrix DOCX", source_type="docx", source=str(path))
        from app.pipeline import run_file

        with _patch_compress(unit):
            result = run_file(path, dest_dir=cards, index=False)
        row.card_path = str(result.markdown_path)
        anim = express_animate_fast(result.markdown_path, work / "express_docx")
        row.gif_path = str(anim.gif_path)
        assert anim.gif_path.is_file()
        row.status = "pass"
    except Exception as exc:  # noqa: BLE001
        row.status = "fail"
        row.error = str(exc)
        row.traceback = traceback.format_exc()
    return row


def run_row_image(work: Path, packages: Path) -> RowResult:
    row = RowResult(
        source_id="image",
        acquire="ingest_image(mock OCR)",
        settle="run_image→card",
        express="animate_fast(pillow)",
    )
    try:
        path = _write_minimal_png(work / "sample.png")
        cards = work / "cards"
        cards.mkdir(parents=True, exist_ok=True)
        unit = fake_unit(title="Matrix Image", source_type="image", source=str(path))
        from app.pipeline import run_image

        src = IngestedSource(
            source_type="image",
            title="Matrix Image",
            text="OCR text: KF matrix force mass acceleration.",
            path=str(path),
            metadata={"ocr_engine": "mock"},
        )
        with (
            patch("app.pipeline.ingest_image", return_value=src),
            _patch_compress(unit),
        ):
            result = run_image(path, dest_dir=cards, index=False)
        row.card_path = str(result.markdown_path)
        anim = express_animate_fast(result.markdown_path, work / "express_image")
        row.gif_path = str(anim.gif_path)
        assert anim.gif_path.is_file()
        row.status = "pass"
    except Exception as exc:  # noqa: BLE001
        row.status = "fail"
        row.error = str(exc)
        row.traceback = traceback.format_exc()
    return row


def run_row_audio(work: Path, packages: Path) -> RowResult:
    row = RowResult(
        source_id="audio",
        acquire="ingest_audio(mock ASR)",
        settle="run_audio→card",
        express="ko_narrate_mock",
    )
    try:
        path = _write_minimal_wav(work / "talk.wav")
        cards = work / "cards"
        cards.mkdir(parents=True, exist_ok=True)
        unit = fake_unit(title="Matrix Audio", source_type="audio", source=str(path))
        from app.pipeline import run_audio

        src = IngestedSource(
            source_type="audio",
            title="Matrix Audio",
            text="这是美债测试录音。收益率上升。",
            path=str(path),
            metadata={"transcript_source": "whisper", "asr": "mock"},
        )
        with (
            patch("app.pipeline.ingest_audio", return_value=src),
            _patch_compress(unit),
        ):
            result = run_audio(path, dest_dir=cards, index=False)
        row.card_path = str(result.markdown_path)

        fake_wav = work / "mock_narration.wav"
        fake_wav.write_bytes(b"RIFF$$$$WAVEfmt " + b"\x00" * 600)

        def _fake_narrate(obj, *, dest_dir, voice_name=None):
            from app.expression.derive import derive_audio_from_ko
            from app.expression.from_ko import AudioRenderResult, write_audio_expression

            expr = derive_audio_from_ko(obj)
            dest_dir = Path(dest_dir)
            dest_dir.mkdir(parents=True, exist_ok=True)
            expr_path = write_audio_expression(expr, dest_dir / "audio_expression.json")
            out = dest_dir / "narration.wav"
            out.write_bytes(fake_wav.read_bytes())
            return AudioRenderResult(expression=expr, expression_path=expr_path, wav_path=out)

        with patch("app.harness.pipeline.narrate_from_ko", side_effect=_fake_narrate):
            with patch("app.harness.registry.require_for_steps", return_value=[]):
                compiled = settle_from_card(
                    result.markdown_path,
                    packages=packages,
                    animate=False,
                    narrate=True,
                )
        row.package_dir = str(compiled.package_dir)
        assert compiled.narration_wav and compiled.narration_wav.is_file()
        row.detail = f"wav={compiled.narration_wav}"
        row.status = "pass"
    except Exception as exc:  # noqa: BLE001
        row.status = "fail"
        row.error = str(exc)
        row.traceback = traceback.format_exc()
    return row


def _run_url_source(
    *,
    source_id: str,
    run_fn_name: str,
    ingest_patch: str,
    url: str,
    source_type: str,
    work: Path,
    packages: Path,
    express: ExpressKind,
) -> RowResult:
    row = RowResult(
        source_id=source_id,
        acquire=f"{run_fn_name}(mock ingest)",
        settle="card",
        express=express,
    )
    try:
        cards = work / "cards"
        cards.mkdir(parents=True, exist_ok=True)
        unit = fake_unit(
            title=f"Matrix {source_id}",
            source_type=source_type,
            source=url,
            url=url,
        )
        src = IngestedSource(
            source_type=source_type,  # type: ignore[arg-type]
            title=f"Matrix {source_id}",
            text=f"Mock transcript for {source_id} integration matrix.",
            url=url,
            metadata={"mock": True, "video_id": f"matrix_{source_id}"},
        )
        import app.pipeline as pipeline
        from app import config as cfg

        # URL runners lack dest_dir — isolate writes under work/knowledge.
        know = work / "knowledge"
        know.mkdir(parents=True, exist_ok=True)
        run_fn = getattr(pipeline, run_fn_name)
        with (
            patch(ingest_patch, return_value=src),
            _patch_compress(unit),
            patch.object(cfg, "KNOWLEDGE_DIR", know),
        ):
            result = run_fn(url, index=False)
        row.card_path = str(result.markdown_path)
        if express == "animate_fast":
            anim = express_animate_fast(result.markdown_path, work / f"express_{source_id}")
            row.gif_path = str(anim.gif_path)
            assert anim.gif_path.is_file()
        else:
            compiled = settle_from_card(
                result.markdown_path, packages=packages, animate=True, narrate=False
            )
            row.package_dir = str(compiled.package_dir)
            assert compiled.animation_gif and compiled.animation_gif.is_file()
            row.gif_path = str(compiled.animation_gif)
        row.status = "pass"
    except Exception as exc:  # noqa: BLE001
        row.status = "fail"
        row.error = str(exc)
        row.traceback = traceback.format_exc()
    return row


def run_row_youtube(work: Path, packages: Path) -> RowResult:
    return _run_url_source(
        source_id="youtube",
        run_fn_name="run_youtube",
        ingest_patch="app.pipeline.ingest_youtube",
        url="https://www.youtube.com/watch?v=matrixyt01",
        source_type="youtube",
        work=work,
        packages=packages,
        express="animate_fast",
    )


def run_row_bilibili(work: Path, packages: Path) -> RowResult:
    return _run_url_source(
        source_id="bilibili",
        run_fn_name="run_bilibili",
        ingest_patch="app.pipeline.ingest_bilibili",
        url="https://www.bilibili.com/video/BVmatrix01",
        source_type="bilibili",
        work=work,
        packages=packages,
        express="ko_animate",
    )


def run_row_twitter(work: Path, packages: Path) -> RowResult:
    return _run_url_source(
        source_id="twitter",
        run_fn_name="run_twitter",
        ingest_patch="app.pipeline.ingest_twitter",
        url="https://x.com/user/status/1234567890",
        source_type="twitter",
        work=work,
        packages=packages,
        express="animate_fast",
    )


def run_row_search(work: Path, packages: Path) -> RowResult:
    row = RowResult(
        source_id="search",
        acquire="search_files dry + one hit settle",
        settle="run_file on hit",
        express="animate_fast(pillow)",
    )
    try:
        root = work / "search_root"
        cards = work / "cards"
        root.mkdir()
        cards.mkdir(parents=True, exist_ok=True)
        hit = root / "hit_gold_note.txt"
        hit.write_text("搜索命中：黄金与美元关系。", encoding="utf-8")
        from app.ingest.search import search_files
        from app.pipeline import run_file

        hits = search_files([root], keyword="gold", extensions=(".txt",))
        assert hits, "search should find hit_gold_note.txt"
        unit = fake_unit(title="Matrix Search Hit", source_type="txt", source=str(hit))
        with _patch_compress(unit):
            result = run_file(hit, dest_dir=cards, index=False)
        row.card_path = str(result.markdown_path)
        anim = express_animate_fast(result.markdown_path, work / "express_search")
        row.gif_path = str(anim.gif_path)
        row.detail = f"hits={len(hits)}"
        assert anim.gif_path.is_file()
        row.status = "pass"
    except Exception as exc:  # noqa: BLE001
        row.status = "fail"
        row.error = str(exc)
        row.traceback = traceback.format_exc()
    return row


def run_row_ecosystem(work: Path, packages: Path) -> RowResult:
    row = RowResult(
        source_id="ecosystem",
        acquire="discover_design_docs + dry_run gate",
        settle="run_file on design doc (compress mocked)",
        express="animate_fast(pillow)",
    )
    try:
        root = work / "eco_root"
        cards = work / "cards"
        root.mkdir()
        cards.mkdir(parents=True, exist_ok=True)
        doc = root / "DESIGN_FACTOR.md"
        doc.write_text("# Factor design\n\nAlpha is residual return.", encoding="utf-8")
        from app.ingest.ecosystem import discover_design_docs, run_ecosystem_ingest
        from app.pipeline import run_file

        hits = discover_design_docs([root], project="factorlib")
        dry = run_ecosystem_ingest("factorlib", [root], dry_run=True, index=False)
        unit = fake_unit(title="Matrix Ecosystem", source_type="md", source=str(doc))
        with _patch_compress(unit):
            result = run_file(doc, dest_dir=cards, index=False)
        row.card_path = str(result.markdown_path)
        anim = express_animate_fast(result.markdown_path, work / "express_eco")
        row.gif_path = str(anim.gif_path)
        row.detail = f"discover_hits={len(hits)} dry_hits={len(dry.hits)}"
        assert anim.gif_path.is_file()
        row.status = "pass"
    except Exception as exc:  # noqa: BLE001
        row.status = "fail"
        row.error = str(exc)
        row.traceback = traceback.format_exc()
    return row


def run_row_setv_snapshot(work: Path, packages: Path) -> RowResult:
    row = RowResult(
        source_id="setv_snapshot",
        acquire="existing cite card",
        settle="compile --from-card",
        express="animate_fast(pillow)",
    )
    try:
        from app import config

        snaps = sorted(
            (config.ROOT / "data" / "knowledge" / "restricted" / "setv" / "snapshots").glob(
                "snapshot_setv_inst_aapl_*.md"
            )
        )
        if not snaps:
            row.status = "skip"
            row.detail = "no AAPL snapshot cite cards under knowledge/restricted"
            return row
        card = snaps[0]
        row.card_path = str(card)
        compiled = settle_from_card(card, packages=packages, animate=False, narrate=False)
        row.package_dir = str(compiled.package_dir)
        anim = express_animate_fast(card, work / "express_setv")
        row.gif_path = str(anim.gif_path)
        assert anim.gif_path.is_file()
        row.status = "pass"
        row.detail = f"card={card.name}"
    except Exception as exc:  # noqa: BLE001
        row.status = "fail"
        row.error = str(exc)
        row.traceback = traceback.format_exc()
    return row


def run_row_setv_family(work: Path, packages: Path) -> RowResult:
    row = RowResult(
        source_id="setv_family",
        acquire="family_view API resolve",
        settle="N/A (read-only multi-card)",
        express="family_view",
    )
    try:
        from app.ui.family_view import resolve_family_view

        view = resolve_family_view(
            "SETV-FAM-AAPL-TV-2024-WDH4",
            lane="proprietary",
            limit=8,
        )
        assert view.get("ok") is True
        assert view.get("members")
        row.detail = f"members={len(view['members'])}"
        row.status = "pass"
    except Exception as exc:  # noqa: BLE001
        row.status = "fail"
        row.error = str(exc)
        row.traceback = traceback.format_exc()
    return row


MATRIX_RUNNERS: list[tuple[str, Callable[[Path, Path], RowResult]]] = [
    ("txt", run_row_txt),
    ("md", run_row_md),
    ("pdf", run_row_pdf),
    ("docx", run_row_docx),
    ("image", run_row_image),
    ("audio", run_row_audio),
    ("youtube", run_row_youtube),
    ("bilibili", run_row_bilibili),
    ("twitter", run_row_twitter),
    ("search", run_row_search),
    ("ecosystem", run_row_ecosystem),
    ("setv_snapshot", run_row_setv_snapshot),
    ("setv_family", run_row_setv_family),
]


def run_full_matrix(root: Path) -> list[RowResult]:
    packages = root / "packages"
    packages.mkdir(parents=True, exist_ok=True)
    results: list[RowResult] = []
    for source_id, runner in MATRIX_RUNNERS:
        work = root / source_id
        work.mkdir(parents=True, exist_ok=True)
        results.append(runner(work, packages))
    return results


def write_evidence(results: list[RowResult], dest: Path) -> Path:
    dest.parent.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    passed = sum(1 for r in results if r.status == "pass")
    failed = [r for r in results if r.status == "fail"]
    skipped = [r for r in results if r.status == "skip"]

    lines = [
        "# Integration Source Matrix — settle → express (non-Cartesian)",
        "",
        "```yaml",
        f"as_of: {stamp}",
        "rule: one express path per signal source (no Cartesian product)",
        f"pass: {passed}/{len(results)}",
        f"fail: {len(failed)}",
        f"skip: {len(skipped)}",
        "```",
        "",
        "## Matrix",
        "",
        "| Source | Acquire | Settle | Express | Status | Detail |",
        "|--------|---------|--------|---------|--------|--------|",
    ]
    for r in results:
        detail = (r.detail or r.error or "").replace("|", "/")[:120]
        lines.append(
            f"| `{r.source_id}` | {r.acquire} | {r.settle} | {r.express} | "
            f"**{r.status}** | {detail} |"
        )

    lines += ["", "## Failures (for fix loop)", ""]
    if not failed:
        lines.append("_None._")
    else:
        for r in failed:
            lines += [
                f"### `{r.source_id}`",
                "",
                f"- error: `{r.error}`",
                "",
                "```",
                (r.traceback or "")[:4000],
                "```",
                "",
            ]

    lines += [
        "",
        "## Notes",
        "",
        "- Network ingest (youtube/bilibili/twitter) mocked at ingest boundary.",
        "- LLM compress mocked; OCR/ASR mocked.",
        "- Audio narrate uses mocked TTS wav (real TTS remains behind `KF_RUN_SLOW`).",
        "- SETV snapshot uses live cite card when present.",
        "",
    ]
    dest.write_text("\n".join(lines) + "\n", encoding="utf-8")

    json_path = dest.with_suffix(".json")
    json_path.write_text(
        json.dumps([asdict(r) for r in results], ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return dest
