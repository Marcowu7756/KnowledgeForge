from __future__ import annotations

from pathlib import Path

from app.ingest.errors import FileIngestError
from app.ingest.pdf import ingest_pdf
from app.models import IngestedSource, SourceType

_TEXT_TYPES: dict[str, SourceType] = {
    ".md": "md",
    ".txt": "txt",
    ".markdown": "md",
}


def _read_text(path: Path) -> str:
    raw = path.read_bytes()
    if not raw.strip():
        raise FileIngestError(f"empty file: {path}")
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise FileIngestError(f"could not decode text file: {path}")


def _ingest_docx(file_path: Path) -> IngestedSource:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise FileIngestError("python-docx is not installed") from exc

    document = Document(str(file_path))
    parts: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n\n".join(parts).strip()
    if not text:
        raise FileIngestError(f"no extractable text in DOCX: {file_path}")

    return IngestedSource(
        source_type="docx",
        title=file_path.stem,
        text=text,
        path=str(file_path),
        metadata={
            "bytes": file_path.stat().st_size,
            "suffix": ".docx",
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
            "read_only_source": True,
        },
    )


def ingest_file(path: str | Path) -> IngestedSource:
    """Read a local document. Source path is recorded; never modified."""
    file_path = Path(path).expanduser().resolve()
    if not file_path.is_file():
        raise FileIngestError(f"not a file: {file_path}")

    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return ingest_pdf(file_path)
    if suffix == ".docx":
        return _ingest_docx(file_path)
    if suffix not in _TEXT_TYPES:
        raise FileIngestError(
            f"unsupported extension {suffix!r}; "
            f"expected {tuple(_TEXT_TYPES) + ('.pdf', '.docx')}"
        )

    text = _read_text(file_path)
    source_type = _TEXT_TYPES[suffix]
    return IngestedSource(
        source_type=source_type,
        title=file_path.stem,
        text=text,
        path=str(file_path),
        metadata={
            "bytes": file_path.stat().st_size,
            "suffix": suffix,
            "read_only_source": True,
        },
    )
